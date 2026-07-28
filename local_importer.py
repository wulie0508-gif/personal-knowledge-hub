#!/usr/bin/env python3
"""Import local documents into the shared Obsidian knowledge network."""

from __future__ import annotations

import hashlib
import html as html_std
import json
import re
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

import codex_curation_queue
import knowledge_pipeline
import knowledge_schema
import knowledge_value
import runtime_config


VAULT = runtime_config.configured_vault()
LOCAL_ROOT = VAULT / "10_Sources" / "Local"
ARTICLE_ROOT = LOCAL_ROOT / "Articles"
SUPPORTED_SUFFIXES = {".md", ".txt", ".html", ".htm", ".json", ".docx", ".pdf"}
SKIP_DIRS = {".git", ".obsidian", "node_modules", "__pycache__"}
MAX_FILES = 500
MAX_TEXT_CHARS = 300_000


@dataclass
class ImportedLocalFile:
    source_path: str
    title: str
    note_path: str = ""
    status: str = "pending"
    text_length: int = 0
    category: str = "其他"
    knowledge_type: str = "参考资料"
    knowledge_value_score: int = 0
    knowledge_priority: str = "参考"
    error: str = ""


def now_text() -> str:
    return datetime.now().astimezone().isoformat(timespec="seconds")


def safe_name(value: str, limit: int = 80) -> str:
    return knowledge_pipeline.safe_name(value, limit)


def source_id(path: Path) -> str:
    return hashlib.sha1(str(path.resolve()).lower().encode("utf-8")).hexdigest()[:12]


def read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def strip_frontmatter(text: str) -> str:
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) == 3:
            return parts[2].lstrip()
    return text


def json_lines(value: Any, level: int = 2) -> list[str]:
    rows: list[str] = []
    if isinstance(value, dict):
        for key, child in value.items():
            rows.append(f"{'#' * min(level, 5)} {key}")
            rows.extend(json_lines(child, level + 1))
    elif isinstance(value, list):
        for child in value:
            if isinstance(child, (dict, list)):
                rows.extend(json_lines(child, level))
            else:
                rows.append(f"- {child}")
    elif value not in (None, ""):
        rows.append(str(value))
    return rows


def extract_file(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    title = path.stem
    if suffix in {".md", ".txt"}:
        text = read_text_file(path)
        if suffix == ".md":
            title_match = re.search(r'(?m)^title:\s*["\']?([^"\'\n]+)', text)
            heading_match = re.search(r"(?m)^#\s+(.+)$", text)
            title = (
                title_match.group(1).strip()
                if title_match
                else heading_match.group(1).strip()
                if heading_match
                else title
            )
            text = strip_frontmatter(text)
        return title, text
    if suffix in {".html", ".htm"}:
        soup = BeautifulSoup(read_text_file(path), "lxml")
        for node in soup(["script", "style", "noscript", "svg"]):
            node.decompose()
        title = (
            (soup.title.get_text(" ", strip=True) if soup.title else "")
            or title
        )
        return title, soup.get_text("\n", strip=True)
    if suffix == ".json":
        value = json.loads(read_text_file(path))
        if isinstance(value, dict):
            title = str(value.get("title") or value.get("name") or title)
        return title, "\n\n".join(json_lines(value))
    if suffix == ".docx":
        document = Document(path)
        title = document.core_properties.title or title
        rows = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    rows.append(" | ".join(values))
        return title, "\n\n".join(rows)
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        title = str((reader.metadata or {}).get("/Title") or title)
        pages = []
        for index, page in enumerate(reader.pages, 1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"## 第 {index} 页\n\n{text}")
        return title, "\n\n".join(pages)
    raise ValueError(f"暂不支持 {suffix or '无扩展名'} 文件")


def iter_source_files(source: Path) -> Iterable[Path]:
    resolved = source.resolve()
    if not resolved.exists():
        raise FileNotFoundError(f"路径不存在：{resolved}")
    if VAULT.resolve() == resolved:
        raise ValueError("不能把当前 Obsidian Vault 再次导入自身")
    if resolved.is_file():
        if resolved.suffix.lower() not in SUPPORTED_SUFFIXES:
            raise ValueError(f"暂不支持 {resolved.suffix or '无扩展名'} 文件")
        yield resolved
        return
    count = 0
    for path in resolved.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        if any(part in SKIP_DIRS for part in path.parts):
            continue
        count += 1
        if count > MAX_FILES:
            raise ValueError(f"单次最多导入 {MAX_FILES} 个文件")
        yield path


def initial_quality(text: str) -> int:
    length = len(re.sub(r"\s+", "", text))
    if length >= 1500:
        return 78
    if length >= 500:
        return 70
    if length >= 120:
        return 60
    if length:
        return 42
    return 15


def output_folder(priority: str, category: str) -> Path:
    if priority == "重点":
        return ARTICLE_ROOT / "重点知识" / safe_name(category, 30)
    if priority == "速览":
        return ARTICLE_ROOT / "资讯速览"
    if priority == "回收建议":
        return ARTICLE_ROOT / "低价值待清理"
    return ARTICLE_ROOT / safe_name(category, 30)


def build_note(
    *,
    source: Path,
    title: str,
    text: str,
    quality: int,
    category: str,
    assessment: knowledge_value.ValueAssessment,
    corpus_namespace: str = knowledge_schema.PERSONAL_MEMORY,
) -> str:
    source_uri = source.resolve().as_uri()
    identity = knowledge_schema.identity_metadata(
        namespace=corpus_namespace,
        fields={},
    )
    rows = [
        "---",
        f"title: {json.dumps(title, ensure_ascii=False)}",
        f"sourceUrl: {json.dumps(source_uri, ensure_ascii=False)}",
        f"source_path: {json.dumps(str(source.resolve()), ensure_ascii=False)}",
        'platform: "local"',
        f"corpus_namespace: {json.dumps(corpus_namespace, ensure_ascii=False)}",
        f"authorship: {json.dumps(identity['authorship'], ensure_ascii=False)}",
        f"confidentiality: {json.dumps(identity['confidentiality'], ensure_ascii=False)}",
        f"engagement_status: {json.dumps(identity['engagement_status'], ensure_ascii=False)}",
        f"stance: {json.dumps(identity['stance'], ensure_ascii=False)}",
        f"persona_influence: {identity['persona_influence']}",
        f"account: {json.dumps(source.parent.name or '本地文件', ensure_ascii=False)}",
        f"category: {json.dumps(category, ensure_ascii=False)}",
        f"quality_score: {quality}",
        f"quality_tier: {json.dumps('高' if quality >= 75 else '中' if quality >= 50 else '低', ensure_ascii=False)}",
        "quality_flags: []",
        f"knowledge_type: {json.dumps(assessment.knowledge_type, ensure_ascii=False)}",
        f"knowledge_value_score: {assessment.value_score}",
        f"knowledge_priority: {json.dumps(assessment.priority, ensure_ascii=False)}",
        f"value_reasons: {json.dumps(assessment.reasons, ensure_ascii=False)}",
        'mastery_status: "未学习"',
        f"content_status: {json.dumps('complete' if text.strip() else 'metadata_only', ensure_ascii=False)}",
        f"text_length: {len(re.sub(r'\\s+', '', text))}",
        f"file_type: {json.dumps(source.suffix.lower().lstrip('.'), ensure_ascii=False)}",
        f"imported_at: {json.dumps(now_text(), ensure_ascii=False)}",
        "---",
        "",
        f"# {title}",
        "",
        f"- 来源：`{source.resolve()}`",
        f"- 分类：{category}",
        f"- 知识价值：{assessment.knowledge_type} · {assessment.priority}（{assessment.value_score}/100）",
        "",
        "<!-- knowledge-value:start -->",
    ]
    if assessment.priority == "重点":
        rows += [
            "> [!important] 重点知识",
            "> 这份本地资料包含可复用的观点、方法或案例。",
            "",
        ]
    if assessment.summary:
        rows += ["## 知识提炼", "", assessment.summary, ""]
    if assessment.highlights:
        rows += ["### 关键点", ""]
        rows += [f"- {item}" for item in assessment.highlights]
        rows.append("")
    rows += [
        "<!-- knowledge-value:end -->",
        "",
        "## 正文",
        "",
        text.strip() or "> 文件中没有提取到可检索文字。",
        "",
    ]
    return "\n".join(rows)


def import_file(
    path: Path,
    corpus_namespace: str = knowledge_schema.PERSONAL_MEMORY,
) -> ImportedLocalFile:
    item = ImportedLocalFile(source_path=str(path.resolve()), title=path.stem)
    try:
        title, text = extract_file(path)
        title = safe_name(html_std.unescape(title), 100)
        text = text[:MAX_TEXT_CHARS].strip()
        quality = initial_quality(text)
        category = knowledge_pipeline.categorize(title, text)
        assessment = knowledge_value.assess(
            title=title,
            body=text,
            quality_score=quality,
            content_status="complete" if text else "metadata_only",
        )
        folder = output_folder(assessment.priority, category)
        folder.mkdir(parents=True, exist_ok=True)
        note = folder / f"{safe_name(title)}--{source_id(path)}.md"
        for previous in ARTICLE_ROOT.rglob(f"*--{source_id(path)}.md"):
            if previous.resolve() != note.resolve():
                previous.unlink()
        note.write_text(
            build_note(
                source=path,
                title=title,
                text=text,
                quality=quality,
                category=category,
                assessment=assessment,
                corpus_namespace=corpus_namespace,
            ).rstrip()
            + "\n",
            encoding="utf-8",
        )
        codex_curation_queue.enqueue(note)
        item.title = title
        item.note_path = str(note)
        item.status = "complete"
        item.text_length = len(re.sub(r"\s+", "", text))
        item.category = category
        item.knowledge_type = assessment.knowledge_type
        item.knowledge_value_score = assessment.value_score
        item.knowledge_priority = assessment.priority
    except Exception as exc:
        item.status = "failed"
        item.error = str(exc)
    return item


def import_path(
    value: str,
    corpus_namespace: str = knowledge_schema.PERSONAL_MEMORY,
) -> dict[str, Any]:
    if corpus_namespace not in knowledge_schema.NAMESPACES - {knowledge_schema.SOURCE_ARCHIVE}:
        raise ValueError(f"unsupported corpus namespace: {corpus_namespace}")
    source = Path(value.strip().strip('"')).expanduser()
    files = list(iter_source_files(source))
    results = [import_file(path, corpus_namespace) for path in files]
    completed = [item for item in results if item.status == "complete"]
    failed = [item for item in results if item.status == "failed"]
    return {
        "message": f"已导入 {len(completed)} 份本地资料，{len(failed)} 份失败",
        "title": f"本地导入：{source.name or source}",
        "output_file": completed[0].note_path if len(completed) == 1 else str(ARTICLE_ROOT),
        "details": {
            "source_path": str(source.resolve()),
            "file_count": len(files),
            "imported_count": len(completed),
            "failed_count": len(failed),
            "items": [asdict(item) for item in results],
        },
    }
